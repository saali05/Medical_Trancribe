# transcription/views.py  —  COMPLETE FILE (replace your existing one)
import traceback

from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login
import base64
from django.core.files.base import ContentFile
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse

from .exports import generate_pdf
from docx import Document

from django.http import HttpResponse

from reportlab.pdfgen import canvas


from .models import Transcription

from .models import AudioRecord, Transcript, ClinicalDocument
from .ai_pipeline import transcribe_audio, process_nlp, structure_clinical_document
# REPLACE with this

# gpt
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.models import User
from django.shortcuts import get_object_or_404

@staff_member_required
def admin_dashboard(request):

    total_users = User.objects.count()

    total_records = AudioRecord.objects.count()

    total_documents = ClinicalDocument.objects.count()

    approved_documents = ClinicalDocument.objects.filter(
        is_approved=True
    ).count()

    recent_records = AudioRecord.objects.order_by(
        '-uploaded_at'
    )[:10]

    context = {
        'total_users': total_users,
        'total_records': total_records,
        'total_documents': total_documents,
        'approved_documents': approved_documents,
        'recent_records': recent_records,
    }

    return render(
        request,
        'transcription/admin_dashboard.html',
        context
    )
# ══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

@login_required
def dashboard(request):
    records = (
    AudioRecord.objects
    .filter(doctor=request.user)
    .select_related('transcript__clinicaldocument')
    .order_by('-uploaded_at')
    )
    return render(request, 'transcription/dashboard.html', {'records': records})


# ══════════════════════════════════════════════════════════════════════════════
#  UPLOAD / RECORD AUDIO  →  runs the full AI pipeline
# ══════════════════════════════════════════════════════════════════════════════

@login_required
def upload_audio(request):

    if request.method == 'POST':

        patient_name = request.POST.get(
            'patient_name',
            ''
        ).strip()

        # NORMAL FILE UPLOAD
        audio_file = request.FILES.get(
            'audio_file'
        )

        # LIVE MICROPHONE RECORDING
        recorded_audio = request.POST.get(
            'recorded_audio'
        )

        # HANDLE RECORDED AUDIO
        if recorded_audio:

            import base64

            from django.core.files.base import ContentFile

            format, audio_string = recorded_audio.split(
                ';base64,'
            )

            ext = format.split('/')[-1]

            audio_file = ContentFile(
                base64.b64decode(audio_string),
                name='recorded_audio.' + ext
            )

        # VALIDATION
        if not audio_file:

            messages.error(
                request,
                'No audio file received. Please try again.'
            )

            return render(
                request,
                'transcription/upload.html'
            )

        # SAVE AUDIO RECORD
        record = AudioRecord.objects.create(

            doctor=request.user,

            audio_file=audio_file,

        )

        try:
            # 2. Whisper ASR — speech → raw text
            audio_path = record.audio_file.path
            raw_text   = transcribe_audio(audio_path)

            # 3. NLP processing — clean + entity extraction
            processed_text, entities = process_nlp(raw_text)

            # 4. Save transcript
            transcript = Transcript.objects.create(
                audio=record,
                raw_text=raw_text,
                processed_text=processed_text,
            )

            # 5. Clinical structuring
            sections = structure_clinical_document(processed_text, entities)

            # 6. Save clinical document
            clinical_doc = ClinicalDocument.objects.create(
                transcript=transcript,
                patient_name=patient_name,
                **sections,
            )

            messages.success(request, 'Transcription complete! Review and edit the document below.')
            return redirect('view_document', pk=clinical_doc.pk)

        except Exception as e:
            # Clean up the orphaned audio record on failure
            record.delete()
            messages.error(request, f'AI pipeline error: {str(e)}. Please try again.')
            return render(request, 'transcription/upload.html')

    return render(request, 'transcription/upload.html')


# ══════════════════════════════════════════════════════════════════════════════
#  VIEW / EDIT CLINICAL DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

@login_required
def view_document(request, pk):
    doc = get_object_or_404(
        ClinicalDocument,
        pk=pk,
        transcript__audio__doctor=request.user,
    )

    if request.method == 'POST':
        # Update all editable fields
        editable_fields = [
            'patient_name', 'chief_complaint', 'history',
            'examination', 'diagnosis', 'prescription',
            'treatment_plan', 'follow_up',
        ]
        for field in editable_fields:
            setattr(doc, field, request.POST.get(field, ''))

        # Approve only if the approve button was clicked
        if 'approve' in request.POST:
            doc.is_approved = True
            messages.success(request, 'Document approved and finalised.')
        else:
            messages.success(request, 'Changes saved successfully.')

        doc.save()
        return redirect('view_document', pk=pk)

    return render(request, 'transcription/document.html', {'doc': doc})


# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT — PDF
# ══════════════════════════════════════════════════════════════════════════════

@login_required
# In views.py, update the call:
def export_pdf(request, pk):
    transcription = Transcription.objects.filter(pk=pk).first()
    if transcription:
        try:
            from .exports import export_transcription_pdf
            return export_transcription_pdf(transcription, request.user)  # ← add request.user
        except Exception as e:
            traceback.print_exc()
            print("TRANSCRIPTION PDF ERROR:", e)

    try:
        doc = ClinicalDocument.objects.get(
            pk=pk,
            transcript__audio__doctor=request.user,
        )
        pdf_bytes = generate_pdf(doc)
        patient  = (doc.patient_name or 'patient').replace(' ', '_').lower()
        date_str = doc.created_at.strftime('%Y%m%d')
        filename = f'clinical_{patient}_{date_str}_DOCAI{doc.pk:05d}.pdf'
        from django.http import HttpResponse
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        traceback.print_exc()
        print("CLINICAL PDF ERROR:", e)
        messages.error(request, f'PDF export failed: {e}')
        return redirect('view_document', pk=pk)


# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT — DOCX
# ══════════════════════════════════════════════════════════════════════════════

@login_required
def export_docx(request, pk):
    """Generate and download a professional clinical Word document."""
    doc = get_object_or_404(
        ClinicalDocument,
        pk=pk,
        transcript__audio__doctor=request.user,
    )

    try:
        docx_bytes = generate_docx(doc)
    except Exception as e:
        messages.error(request, f'DOCX export failed: {e}')
        return redirect('view_document', pk=pk)

    patient  = (doc.patient_name or 'patient').replace(' ', '_').lower()
    date_str = doc.created_at.strftime('%Y%m%d')
    filename = f'clinical_{patient}_{date_str}_DOCAI{doc.pk:05d}.docx'

    response = HttpResponse(
        docx_bytes,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def register_view(request):

    if request.method == 'POST':

        form = UserCreationForm(request.POST)

        if form.is_valid():

            user = form.save()

            login(request, user)

            return redirect('dashboard')

    else:

        form = UserCreationForm()

    return render(

        request,

        'registration/register.html',

        {

            'form': form

        }

    )

def download_pdf(request, pk):

    from django.http import HttpResponse

    from reportlab.platypus import (

        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle

    )

    from reportlab.lib import colors

    from reportlab.lib.styles import getSampleStyleSheet

    from reportlab.lib.pagesizes import A4

    from reportlab.platypus.flowables import HRFlowable

    from reportlab.lib.enums import TA_CENTER

    from reportlab.lib.styles import ParagraphStyle

    from reportlab.pdfbase.ttfonts import TTFont

    from reportlab.pdfbase import pdfmetrics

    from reportlab.platypus import Image

    from reportlab.lib.units import inch

    from .models import ClinicalDocument

    doc = get_object_or_404(
    ClinicalDocument,
    pk=pk
)

    pdf = SimpleDocTemplate(

        response,

        pagesize=A4,

        rightMargin=40,

        leftMargin=40,

        topMargin=40,

        bottomMargin=30

    )

    elements = []

    styles = getSampleStyleSheet()

    # =====================================================
    # CUSTOM STYLES
    # =====================================================

    title_style = ParagraphStyle(

        'TitleStyle',

        parent=styles['Heading1'],

        fontSize=24,

        leading=30,

        textColor=colors.HexColor("#0f172a"),

        alignment=TA_CENTER,

        spaceAfter=5,

    )

    subtitle_style = ParagraphStyle(

        'SubtitleStyle',

        parent=styles['BodyText'],

        fontSize=11,

        textColor=colors.HexColor("#64748b"),

        alignment=TA_CENTER,

        spaceAfter=20,

    )

    section_style = ParagraphStyle(

        'SectionStyle',

        parent=styles['Heading2'],

        fontSize=15,

        leading=20,

        textColor=colors.white,

        backColor=colors.HexColor("#0f172a"),

        spaceBefore=16,

        spaceAfter=10,

        leftIndent=8,

    )

    body_style = ParagraphStyle(

        'BodyStyle',

        parent=styles['BodyText'],

        fontSize=11,

        leading=20,

        textColor=colors.HexColor("#1e293b"),

    )

    footer_style = ParagraphStyle(

        'FooterStyle',

        parent=styles['BodyText'],

        fontSize=9,

        alignment=TA_CENTER,

        textColor=colors.HexColor("#94a3b8"),

    )

    # =====================================================
    # HEADER
    # =====================================================

    elements.append(

        Paragraph(

            "DOCAI MEDICAL CENTER",

            title_style

        )

    )

    elements.append(

        Paragraph(

            "AI Powered Clinical Documentation System",

            subtitle_style

        )

    )

    elements.append(

        HRFlowable(

            width="100%",

            thickness=1.5,

            color=colors.HexColor("#0ea5a0")

        )

    )

    elements.append(Spacer(1, 18))

    # =====================================================
    # DOCTOR INFO
    # =====================================================

    doctor_data = [

        [

            Paragraph(

                "<b>Doctor:</b> Dr. Salim Rahman",

                body_style

            ),

            Paragraph(

                "<b>Clinic:</b> DocAI Medical Center",

                body_style

            )

        ]

    ]

    doctor_table = Table(

        doctor_data,

        colWidths=[250, 250]

    )

    doctor_table.setStyle(

        TableStyle([

            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),

            ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#cbd5e1")),

            ('BOTTOMPADDING', (0,0), (-1,-1), 10),

            ('TOPPADDING', (0,0), (-1,-1), 10),

        ])

    )

    elements.append(doctor_table)

    elements.append(Spacer(1, 18))

    # =====================================================
    # PATIENT INFO
    # =====================================================

    patient_data = [

        [

            "Patient Name",

            "Age",

            "Mobile",

            "Date"

        ],

        [

            doc.patient_name or "-",

            "39",

            "9876543210",

            doc.created_at.strftime("%d-%m-%Y")

        ]

    ]

    patient_table = Table(

        patient_data,

        colWidths=[170, 80, 140, 110]

    )

    patient_table.setStyle(

        TableStyle([

            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0f172a")),

            ('TEXTCOLOR', (0,0), (-1,0), colors.white),

            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),

            ('FONTSIZE', (0,0), (-1,0), 11),

            ('BOTTOMPADDING', (0,0), (-1,0), 12),

            ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#f8fafc")),

            ('GRID', (0,0), (-1,-1), 1, colors.HexColor("#cbd5e1")),

            ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),

            ('BOTTOMPADDING', (0,1), (-1,-1), 10),

            ('TOPPADDING', (0,1), (-1,-1), 10),

        ])

    )

    elements.append(patient_table)

    elements.append(Spacer(1, 24))

    # =====================================================
    # CLINICAL SECTIONS
    # =====================================================

    sections = [

        ("Chief Complaint", doc.chief_complaint),

        ("Present Illness", doc.history),

        ("Physical Examination", doc.examination),

        ("Diagnosis", doc.diagnosis),

        ("Prescription", doc.prescription),

        ("Treatment Plan", doc.treatment_plan),

        ("Follow Up Instructions", doc.follow_up),

    ]

    for title, content in sections:

        elements.append(

            Paragraph(title, section_style)

        )

        elements.append(

            Spacer(1, 6)

        )

        elements.append(

            Paragraph(

                content or "Not available.",

                body_style

            )

        )

        elements.append(

            Spacer(1, 14)

        )

    # =====================================================
    # FOOTER
    # =====================================================

    elements.append(Spacer(1, 30))

    elements.append(

        HRFlowable(

            width="100%",

            thickness=1,

            color=colors.HexColor("#cbd5e1")

        )

    )

    elements.append(Spacer(1, 10))

    elements.append(

        Paragraph(

            "Generated by DocAI Clinical Intelligence",

            footer_style

        )

    )

    elements.append(

        Paragraph(

            "Confidential Medical Document",

            footer_style

        )

    )

    # BUILD PDF

    pdf.build(elements)

    return response

def download_docx(request, pk):

    doc = get_object_or_404(
    ClinicalDocument,
    pk=pk
)

    # TITLE

    document.add_heading(
        'DocAI Clinical Report',
        level=1
    )

    # CONTENT

    document.add_paragraph(
        f"Patient: {doc.patient_name}"
    )

    document.add_paragraph(
        f"Chief Complaint: {doc.chief_complaint}"
    )

    document.add_paragraph(
        f"History: {doc.history}"
    )

    document.add_paragraph(
        f"Diagnosis: {doc.diagnosis}"
    )

    document.add_paragraph(
        f"Prescription: {doc.prescription}"
    )

    document.add_paragraph(
        f"Treatment Plan: {doc.treatment_plan}"
    )

    document.add_paragraph(
        f"Follow Up: {doc.follow_up}"
    )

    # RESPONSE

    response = HttpResponse(

        content_type=

        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    )

    response['Content-Disposition'] = (

        f'attachment; filename="report_{pk}.docx"'

    )

    document.save(response)

    return response

# medtranscribe/views.py

import os
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Transcription
from .forms import AudioUploadForm
# Import transcriber from the medtranscribe app
from medtranscribe.transcriber import process_audio

def upload_audio(request):
    if request.method == 'POST':
        form = AudioUploadForm(request.POST, request.FILES)
        if form.is_valid():
            transcription = form.save()
            try:
                result = process_audio(transcription.audio_file.path)

                transcription.conversation_json    = result["conversation"]
                med = result["medical_summary"]
                transcription.chief_complaint      = med.get("chief_complaint", "")
                transcription.present_illness      = med.get("present_illness", "")
                transcription.diagnosis            = med.get("diagnosis", "")
                transcription.prescription         = med.get("prescription", "")
                transcription.treatment_plan       = med.get("treatment_plan", "")
                transcription.follow_up            = med.get("follow_up", "")
                transcription.vitals               = med.get("vitals", "")
                transcription.past_medical_history = med.get("past_medical_history", "")
                transcription.save()

                messages.success(request, "Transcription complete!")
                return redirect('transcription_detail', pk=transcription.pk)

            except Exception as e:
                messages.error(request, f"Error: {str(e)}")
                transcription.delete()
    else:
        form = AudioUploadForm()

    return render(request, 'transcription/upload.html', {'form': form})


def transcription_detail(request, pk):
    t = get_object_or_404(Transcription, pk=pk)
    return render(request, 'transcription/detail.html', {'transcription': t})


def transcription_list(request):
    transcriptions = Transcription.objects.order_by('-created_at')
    return render(request, 'transcription/list.html', {'transcriptions': transcriptions})